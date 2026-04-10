from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from .errors import AppError
from .schemas import DocumentParseResult, SectionNode

SPECIAL_TITLE_MAP = {
    "摘要": "abstract_cn",
    "中文摘要": "abstract_cn",
    "abstract": "abstract_en",
    "参考文献": "references",
    "references": "references",
    "致谢": "acknowledgements",
    "致謝": "acknowledgements",
    "acknowledgements": "acknowledgements",
    "acknowledgment": "acknowledgements",
    "appendix": "appendix",
    "附录": "appendix",
}


def normalize_title(value: str) -> str:
    text = value.strip().strip("#").strip()
    return re.sub(r"[\s:：\-\._]+", "", text).lower()


def detect_heading(paragraph_text: str, style_name: str | None) -> tuple[bool, str, int]:
    text = paragraph_text.strip()
    style = (style_name or "").strip().lower()
    if style.startswith("heading"):
        level_text = re.sub(r"\D+", "", style)
        level = int(level_text) if level_text else 1
        return True, text, min(level, 3)

    markdown_match = re.match(r"^(#{1,3})\s+(.+)$", text)
    if markdown_match:
        hashes, title = markdown_match.groups()
        return True, title.strip(), len(hashes)

    normalized = normalize_title(text)
    if normalized in SPECIAL_TITLE_MAP:
        return True, text, 1

    if re.match(r"^第[一二三四五六七八九十百0-9]+章", text):
        return True, text, 1

    return False, "", 0


def split_keywords(text: str, english: bool) -> tuple[str, str]:
    lines = [line.strip() for line in text.splitlines()]
    body_lines: list[str] = []
    keywords = ""
    prefixes = ["keywords", "keyword"] if english else ["关键词", "關鍵詞"]
    for line in lines:
        normalized = line.lower().replace("：", ":")
        matched = False
        for prefix in prefixes:
            if normalized.startswith(prefix):
                keywords = line.split(":", 1)[-1].split("：", 1)[-1].strip()
                matched = True
                break
        if not matched:
            body_lines.append(line)
    return "\n".join(body_lines).strip(), keywords


def parse_docx(path: Path) -> DocumentParseResult:
    if path.suffix.lower() != ".docx":
        raise AppError("DOCX_INVALID", "仅支持 .docx 文件", status_code=400)

    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover - upstream parser detail
        raise AppError("DOCX_INVALID", "无法解析上传的 .docx 文件", details={"reason": str(exc)}, status_code=400) from exc

    sections: list[SectionNode] = []
    warnings: list[str] = []
    current_title = ""
    current_kind = "body"
    current_lines: list[str] = []
    front_matter: list[str] = []

    def flush_current() -> None:
        nonlocal current_title, current_kind, current_lines
        content = "\n".join(line for line in current_lines if line.strip()).strip()
        if current_title or content:
            sections.append(
                SectionNode(
                    kind=current_kind,
                    title=current_title or "正文",
                    content=content,
                )
            )
        current_title = ""
        current_kind = "body"
        current_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            if current_lines:
                current_lines.append("")
            continue

        style_name = paragraph.style.name if paragraph.style is not None else None
        is_heading, heading_title, _level = detect_heading(text, style_name)
        if is_heading:
            flush_current()
            current_title = heading_title
            current_kind = SPECIAL_TITLE_MAP.get(normalize_title(heading_title), "body")
            continue

        if current_title:
            current_lines.append(text)
        else:
            front_matter.append(text)

    flush_current()

    if not sections and front_matter:
        sections.append(SectionNode(kind="body", title="正文", content="\n".join(front_matter).strip()))
    elif front_matter:
        warnings.append("检测到未归类的前置内容，已并入正文开头。")
        for section in sections:
            if section.kind == "body":
                merged = "\n".join(front_matter + [section.content]).strip()
                section.content = merged
                break
        else:
            sections.insert(0, SectionNode(kind="body", title="正文", content="\n".join(front_matter).strip()))

    if not sections:
        raise AppError("CONTENT_EMPTY", "文档中没有可用文本内容", status_code=400)

    return DocumentParseResult(
        source_type="docx",
        metadata={"paragraph_count": str(len(document.paragraphs))},
        sections=sections,
        warnings=warnings,
    )
